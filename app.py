import streamlit as st
from datetime import datetime
from dateutil.relativedelta import relativedelta
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# --- VERİ SETİ (ASGARİ ÜCRET VE KIDEM TAVANI) ---
ASGARI_UCRET_TABLOSU = {
    2026: 25000.00, 2025: 20002.12, 2024: 17002.12, 
    2023: 13414.50, 2022: 6471.00, 2021: 3577.50, 2020: 2943.00, 2019: 2558.40
}
KIDEM_TAVANI_TABLOSU = {
    2026: 60000.00, 2025: 50000.00, 2024: 41828.42, 
    2023: 23489.83, 2022: 15371.40, 2021: 8284.51, 2020: 7117.17
}

def format_tl(value):
    return "{:,.2f} TL".format(value).replace(",", "X").replace(".", ",").replace("X", ".")

def kesinti_hesapla(brut, tip="standart"):
    dv = brut * 0.00759
    if tip == "kıdem":
        return {"brut": brut, "dv": dv, "gv": 0.0, "sgk": 0.0, "issizlik": 0.0, "toplam": dv, "net": brut - dv}
    sgk = brut * 0.14
    issizlik = brut * 0.01
    gv = (brut - (sgk + issizlik)) * 0.15 
    toplam = dv + sgk + gv + issizlik
    return {"brut": brut, "dv": dv, "gv": gv, "sgk": sgk, "issizlik": issizlik, "toplam": toplam, "net": brut - toplam}

def create_docx(isim, sections_list):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)
    
    t = doc.add_heading('BİLİRKİŞİ RAPORU HESAPLAMA DÖKÜMÜ', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Davacı: {isim}")
    doc.add_paragraph(f"Rapor Tarihi: {datetime.now().strftime('%d.%m.%Y')}")

    for title, df in sections_list:
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        table.style = 'Table Grid'
        for j in range(df.shape[1]):
            table.cell(0, j).text = str(df.columns[j])
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                table.cell(i + 1, j).text = str(df.values[i, j])
        doc.add_paragraph("\n")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

st.set_page_config(page_title="Bilirkişi Raporu Pro", layout="wide")
st.markdown("<h2 style='text-align: center;'>⚖️ İşçilik Alacakları Bilirkişi Raporu</h2>", unsafe_allow_html=True)

with st.sidebar:
    st.header("Giriş Parametreleri")
    isim = st.text_input("İşçinin Adı Soyadı", "Onur Erol")
    g_tarih = st.date_input("İşe Giriş Tarihi", datetime(2020, 1, 15))
    c_tarih = st.date_input("İşten Çıkış Tarihi", datetime(2021, 7, 26))
    son_brut = st.number_input("Brüt Ücret (TL)", value=5595.11)
    yemek_ucreti = st.number_input("Yemek Ücreti (Aylık Brüt)", value=350.0)
    fm_saat = st.number_input("Haftalık Fazla Mesai Saati", value=12.0)
    izin_gun = st.number_input("Kalan İzin Günü", value=14)
    ubgt_yillik = st.number_input("Yıllık Çalışılan UBGT Günü", value=5)

if st.button("HESAPLA VE DETAYLI RAPORU OLUŞTUR", type="primary"):
    # --- 1. ÖNCELİKLİ HESAPLAMALAR ---
    delta = relativedelta(c_tarih, g_tarih)
    yil, ay, gun = delta.years, delta.months, delta.days
    toplam_gun = (c_tarih - g_tarih).days
    giydirilmis_brut = son_brut + yemek_ucreti
    tavan = KIDEM_TAVANI_TABLOSU.get(c_tarih.year, 41828.42)
    esas_kidem = min(giydirilmis_brut, tavan)
    d_asgari_cikis = ASGARI_UCRET_TABLOSU.get(c_tarih.year, 3577.50)
    maas_orani = son_brut / d_asgari_cikis
    ihbar_hafta = 2 if toplam_gun < 180 else 4 if toplam_gun < 540 else 6 if toplam_gun < 1080 else 8

    word_sections = []

    # --- 3. ÜCRET TESPİTİ ---
    st.markdown("### 3. Hesaplamalarda Kullanılacak Ücret Miktarları İle İlgili Tespit")
    ucret_data = [
        ["Brüt Ücret", format_tl(son_brut)],
        ["Yemek Ücreti", format_tl(yemek_ucreti)],
        ["Giydirilmiş Brüt Ücret", format_tl(giydirilmis_brut)],
        ["Kıdem Tazminatı Tavan Tutarı", format_tl(tavan)],
        ["Kıdem Tazminatına Esas Ücret", format_tl(esas_kidem)],
        ["İhbar Tazminatına Esas Ücret", format_tl(son_brut)],
        ["Dönem Asgari Ücret (" + str(c_tarih.year) + ")", format_tl(d_asgari_cikis)],
        ["Brüt Ücretin Asgari Ücrete Oranı", "{:.5f}".format(maas_orani).replace(".", ",")]
    ]
    df_ucret = pd.DataFrame(ucret_data, columns=["Kalem", "Miktar"])
    st.table(df_ucret)
    word_sections.append(("3. ÜCRET TESPİTLERİ", df_ucret))

    # --- 4. KIDEM ---
    st.markdown("### 4. Kıdem ve İhbar Tazminatının Hesaplanması")
    st.caption(f"Hizmet Süresi: {yil} Yıl {ay} Ay {gun} Gün")
    b_kidem = (esas_kidem * yil) + (esas_kidem / 12 * ay) + (esas_kidem / 365 * gun)
    k_res = kesinti_hesapla(b_kidem, "kıdem")
    df_kidem = pd.DataFrame([
        [format_tl(esas_kidem), "x", f"{yil} Yıl", "=", format_tl(esas_kidem * yil)],
        [f"{format_tl(esas_kidem)} / 12", "x", f"{ay} Ay", "=", format_tl(esas_kidem/12*ay)],
        [f"{format_tl(esas_kidem)} / 365", "x", f"{gun} Gün", "=", format_tl(esas_kidem/365*gun)],
        ["**TOPLAM BRÜT**", "", "", "", f"**{format_tl(b_kidem)}**"],
        ["Damga Vergisi (Binde 7,59)", "", "", "=", format_tl(k_res['dv'])],
        ["**NET KIDEM**", "", "", "", f"**{format_tl(k_res['net'])}**"]
    ])
    st.table(df_kidem)
    word_sections.append(("4. KIDEM TAZMİNATI", df_kidem))

    # --- 4.1 İHBAR ---
    b_ihbar = (son_brut / 30) * 7 * ihbar_hafta
    i_res = kesinti_hesapla(b_ihbar)
    df_ihbar = pd.DataFrame([
        [f"{format_tl(son_brut)} / 30", "x", f"7 Gün x {ihbar_hafta} Hafta", "=", format_tl(b_ihbar)],
        ["Gelir Vergisi (%15)", "", "", "=", format_tl(i_res['gv'])],
        ["Damga Vergisi (Binde 7,59)", "", "", "=", format_tl(i_res['dv'])],
        ["**NET İHBAR**", "", "", "", f"**{format_tl(i_res['net'])}**"]
    ])
    st.markdown(f"**İhbar Tazminatı Hesabı ({ihbar_hafta} Hafta):**")
    st.table(df_ihbar)
    word_sections.append(("4.1 İHBAR TAZMİNATI", df_ihbar))

    # --- 5. YILLIK İZİN ---
    st.markdown("### 5. Yıllık İzin Ücretinin Hesaplanması")
    b_izin = (son_brut / 30) * izin_gun
    z_res = kesinti_hesapla(b_izin)
    df_izin = pd.DataFrame([
        ["Brüt İzin Alacağı", f"{format_tl(son_brut/30)} x {izin_gun} Gün", "=", format_tl(b_izin)],
        ["SGK Primi (%14)", "", "=", format_tl(z_res['sgk'])],
        ["İşsizlik Sigortası (%1)", "", "=", format_tl(z_res['issizlik'])],
        ["Gelir Vergisi (%15)", "", "=", format_tl(z_res['gv'])],
        ["Damga Vergisi (Binde 7,59)", "", "=", format_tl(z_res['dv'])],
        ["**NET YILLIK İZİN**", "", "=", f"**{format_tl(z_res['net'])}**"]
    ])
    st.table(df_izin)
    word_sections.append(("5. YILLIK İZİN ÜCRETİ", df_izin))

    # --- 6. FAZLA MESAİ ---
    st.markdown("### 6. Fazla Mesai Ücretinin Hesaplanması")
    fm_brut, fm_rows = 0, []
    for y in range(g_tarih.year, c_tarih.year + 1):
        d_maas = ASGARI_UCRET_TABLOSU.get(y, 25000.0) * maas_orani
        sz = (d_maas / 225) * 1.5
        bas = max(g_tarih, datetime(y, 1, 1).date()); bit = min(c_tarih, datetime(y, 12, 31).date())
        h_say = max(0, (bit - bas).days / 7)
        donem_fm = h_say * fm_saat * sz; fm_brut += donem_fm
        fm_rows.append([f"{y}", f"({format_tl(d_maas)}/225)*1,5", f"{fm_saat} Saat * {h_say:.2f} Hafta", format_tl(donem_fm)])
    
    df_fm_d = pd.DataFrame(fm_rows, columns=["Dönem", "Saatlik Zamlı", "FM Süresi", "Brüt"])
    st.table(df_fm_d)
    word_sections.append(("6. FAZLA MESAİ DÖNEMSEL DÖKÜM", df_fm_d))
    
    fm_res = kesinti_hesapla(fm_brut)
    df_fm_k = pd.DataFrame([["SGK Primi (%14)", format_tl(fm_res['sgk'])], ["İşsizlik (%1)", format_tl(fm_res['issizlik'])], ["Gelir Vergisi (%15)", format_tl(fm_res['gv'])], ["Damga V. (Binde 7,59)", format_tl(fm_res['dv'])], ["**Net FM**", format_tl(fm_res['net'])]], columns=["Kesinti Oranı", "Tutar"])
    st.table(df_fm_k)
    word_sections.append(("6.1 FAZLA MESAİ KESİNTİ DÖKÜMÜ", df_fm_k))

    # --- 6.1 UBGT ---
    st.markdown("### 6.1 UBGT Ücretinin Hesaplanması")
    ub_brut, ub_rows = 0, []
    for y in range(g_tarih.year, c_tarih.year + 1):
        d_m = ASGARI_UCRET_TABLOSU.get(y, 25000.0) * maas_orani; gu = d_m / 30
        bas = max(g_tarih, datetime(y, 1, 1).date()); bit = min(c_tarih, datetime(y, 12, 31).date())
        d_g = ubgt_yillik * ((bit - bas).days / 365)
        d_ub = d_g * gu; ub_brut += d_ub
        ub_rows.append([f"{y}", f"{format_tl(d_m)} / 30", f"{d_g:.2f} Gün", format_tl(d_ub)])
    
    df_ub_d = pd.DataFrame(ub_rows, columns=["Dönem", "Günlük Ücret", "Gün Sayısı", "Brüt"])
    st.table(df_ub_d)
    word_sections.append(("6.2 UBGT DÖNEMSEL DÖKÜM", df_ub_d))
    
    u_res = kesinti_hesapla(ub_brut)
    df_ub_k = pd.DataFrame([["SGK Primi (%14)", format_tl(u_res['sgk'])], ["Gelir Vergisi (%15)", format_tl(u_res['gv'])], ["Damga V. (Binde 7,59)", format_tl(u_res['dv'])], ["**Net UBGT**", format_tl(u_res['net'])]], columns=["Kesinti Oranı", "Tutar"])
    st.table(df_ub_k)
    word_sections.append(("6.3 UBGT KESİNTİ DÖKÜMÜ", df_ub_k))

    # --- 7. İCMAL ---
    st.markdown("### 7. Sonuç ve İcmal (Özet) Tablosu")
    g_brut = b_kidem + b_ihbar + b_izin + fm_brut + ub_brut
    g_kes = k_res['toplam'] + i_res['toplam'] + z_res['toplam'] + fm_res['toplam'] + u_res['toplam']
    g_net = k_res['net'] + i_res['net'] + z_res['net'] + fm_res['net'] + u_res['net']
    df_icmal = pd.DataFrame([
        ["Kıdem", format_tl(b_kidem), format_tl(k_res['toplam']), format_tl(k_res['net'])],
        ["İhbar", format_tl(b_ihbar), format_tl(i_res['toplam']), format_tl(i_res['net'])],
        ["İzin", format_tl(b_izin), format_tl(z_res['toplam']), format_tl(z_res['net'])],
        ["FM", format_tl(fm_brut), format_tl(fm_res['toplam']), format_tl(fm_res['net'])],
        ["UBGT", format_tl(ub_brut), format_tl(u_res['toplam']), format_tl(u_res['net'])],
        ["**GENEL TOPLAM**", "**"+format_tl(g_brut)+"**", "**"+format_tl(g_kes)+"**", "**"+format_tl(g_net)+"**"]
    ], columns=["Alacak Kalemi", "Brüt Tutar", "Kesintiler", "Net Ödenecek"])
    st.table(df_icmal)
    word_sections.append(("7. SONUÇ VE İCMAL TABLOSU", df_icmal))

    # WORD DOWNLOAD
    st.download_button(label="📥 Raporu Word Olarak İndir", data=create_docx(isim, word_sections), file_name=f"Bilirkişi_Raporu_{isim}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
